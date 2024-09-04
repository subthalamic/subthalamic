import pandas as pd
from datetime import datetime
import tkinter as tk
from tkinter import ttk
from tkcalendar import DateEntry
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

f = pd.read_csv('formate-scurt.csv')


def toggle(ii):
    if checkbox_val[ii].get():
        param_ap1[ii].config(state='normal')
        param_ap2[ii].config(state='normal')
    else:
        param_ap1[ii].config(state='disabled')
        param_ap2[ii].config(state='disabled')


def disable_entries(event):
    selected_option = combo_operatiune.get()
    if selected_option != 'Analiză':
        entry_nume_proba.config(state='disabled')
        combo_tip_proba.config(state='disabled')
        entry_seria.config(state='disabled')
        entry_analiza.config(state='disabled')
        entry_solicitantCD.config(state='disabled')
        entry_solicitantCC.config(state='disabled')
        entry_data_eliberare.config(state='disabled')
        entry_utilizator.config(state='disabled')
        entry_rezultat.config(state='disabled')

        entry_nume_proba.delete(0, tk.END)
        combo_tip_proba.current(0)
        entry_seria.delete(0, tk.END)
        entry_analiza.delete(0, tk.END)
        entry_data_intrare.set_date(datetime.now())
        entry_data_eliberare.set_date(datetime.now())
        entry_rezultat.delete(0, tk.END)
        entry_observatii.delete(0, tk.END)

    else:
        entry_nume_proba.config(state='normal')
        combo_tip_proba.config(state='readonly')
        entry_seria.config(state='normal')
        entry_analiza.config(state='normal')
        entry_solicitantCD.config(state='normal')
        entry_solicitantCC.config(state='normal')
        entry_data_eliberare.config(state='normal')
        entry_utilizator.config(state='normal')
        entry_rezultat.config(state='normal')

        entry_nume_proba.delete(0, tk.END)
        combo_tip_proba.current(0)
        entry_seria.delete(0, tk.END)
        entry_analiza.delete(0, tk.END)
        entry_data_intrare.set_date(datetime.now())
        entry_data_eliberare.set_date(datetime.now())
        entry_rezultat.delete(0, tk.END)
        entry_observatii.delete(0, tk.END)


def completare_registre(entry):

    r = [pd.read_csv(f'registre/{i}.csv') for i in range(23)]

    if entry.loc[0, 'Operatiune'] == 'Analiză':
        col_type = entry.loc[0, 'Produs']
    else:
        col_type = entry.loc[0, 'Operatiune']

    # Balanta Mettler Toledo AB204-S
    if entry.loc[0, 'aparate'][0]:
        nr0 = len(r[0])
        r[0].loc[nr0, 'Nr. crt.'] = nr0 + 1
        r[0].loc[nr0, 'Data'] = entry.loc[0, 'Data_in']
        r[0].loc[nr0, 'Operaţia/Denumirea probei'] = col_type
        r[0].loc[nr0, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[0].loc[nr0, 'Masa cântărită, g'] = entry.loc[0, 'aparate1'][0]
        r[0].loc[nr0, 'Semnătura persoanei care a efectuat operaţia'] = entry.loc[0, 'Utilizator']
        r[0].loc[nr0, 'Observaţii'] = entry.loc[0, 'Obs']
        r[0].to_csv('registre/0.csv', index=False, header=True)

    # Balanta Radwag – Partner PS4500 R2
    if entry.loc[0, 'aparate'][1]:
        nr1 = len(r[1])
        r[1].loc[nr1, 'Nr. crt.'] = nr1 + 1
        r[1].loc[nr1, 'Data'] = entry.loc[0, 'Data_in']
        r[1].loc[nr1, 'Operaţia/Denumirea probei'] = col_type
        r[1].loc[nr1, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[1].loc[nr1, 'Masa cântărită, g'] = entry.loc[0, 'aparate1'][1]
        r[1].loc[nr1, 'Semnătura persoanei care a efectuat operaţia'] = entry.loc[0, 'Utilizator']
        r[1].loc[nr1, 'Observaţii'] = entry.loc[0, 'Obs']
        r[1].to_csv('registre/1.csv', index=False, header=True)

    # Balanta VWR 611-2271
    if entry.loc[0, 'aparate'][2]:
        nr2 = len(r[2])
        r[2].loc[nr2, 'Nr. crt.'] = nr2 + 1
        r[2].loc[nr2, 'Data'] = entry.loc[0, 'Data_in']
        r[2].loc[nr2, 'Operaţia/Denumirea probei'] = col_type
        r[2].loc[nr2, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[2].loc[nr2, 'Masa cântărită, g'] = entry.loc[0, 'aparate1'][2]
        r[2].loc[nr2, 'Semnătura persoanei care a efectuat operaţia'] = entry.loc[0, 'Utilizator']
        r[2].loc[nr2, 'Observaţii'] = entry.loc[0, 'Obs']
        r[2].to_csv('registre/2.csv', index=False, header=True)

    # Balanta Partner AS 310 R2
    if entry.loc[0, 'aparate'][3]:
        nr3 = len(r[3])
        r[3].loc[nr3, 'Nr. crt.'] = nr3 + 1
        r[3].loc[nr3, 'Data'] = entry.loc[0, 'Data_in']
        r[3].loc[nr3, 'Operaţia/Denumirea probei'] = col_type
        r[3].loc[nr3, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[3].loc[nr3, 'Masa cântărită, g'] = entry.loc[0, 'aparate1'][3]
        r[3].loc[nr3, 'Semnătura persoanei care a efectuat operaţia'] = entry.loc[0, 'Utilizator']
        r[3].loc[nr3, 'Observaţii'] = entry.loc[0, 'Obs']
        r[3].to_csv('registre/3.csv', index=False, header=True)

    # Balanţa termica VWR MB 160
    if entry.loc[0, 'aparate'][4]:
        nr4 = len(r[4])
        r[4].loc[nr4, 'Nr. crt.'] = nr4 + 1
        r[4].loc[nr4, 'Data'] = entry.loc[0, 'Data_in']
        r[4].loc[nr4, 'Operaţia'] = entry.loc[0, 'Operatiune']
        r[4].loc[nr4, 'Denumire probă/Serie'] = f'{entry.loc[0, "Produs"]}/ {entry.loc[0, "Seria"]}'
        r[4].loc[nr4, 'Semnătura utilizator'] = entry.loc[0, 'Utilizator']
        r[4].loc[nr4, 'Observaţii'] = entry.loc[0, 'Obs']
        r[4].to_csv('registre/4.csv', index=False, header=True)

    # Cuptor de calcinare Nabertherm
    if entry.loc[0, 'aparate'][5]:
        nr5 = len(r[5])
        r[5].loc[nr5, 'Nr. crt.'] = nr5 + 1
        r[5].loc[nr5, 'Data'] = entry.loc[0, 'Data_in']
        r[5].loc[nr5, 'Operaţia/Denumirea probei'] = col_type
        r[5].loc[nr5, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[5].loc[nr5, 'Temperatură ºC/ Durata'] = entry.loc[0, 'aparate1'][5]
        r[5].loc[nr5, 'Nume și prenume persoana care a efectuat operația'] = entry.loc[0, 'Utilizator']
        r[5].loc[nr5, 'Observaţii'] = entry.loc[0, 'Obs']
        r[5].to_csv('registre/5.csv', index=False, header=True)

    # Centrifuga Ortoalesa Digicen 21R
    if entry.loc[0, 'aparate'][6]:
        nr6 = len(r[6])
        r[6].loc[nr6, 'Data'] = entry.loc[0, 'Data_in']
        r[6].loc[nr6, 'Denumire probă'] = entry.loc[0, 'Produs']
        r[6].loc[nr6, 'Tipul probei/Lot'] = f'{entry.loc[0, "Tip_produs"]}/ {entry.loc[0, "Seria"]}'
        r[6].loc[nr6, 'Timp de centrifugare'] = entry.loc[0, 'aparate1'][6]
        r[6].loc[nr6, 'RPM'] = entry.loc[0, 'aparate2'][6]
        r[6].loc[nr6, 'Semnătura operator'] = entry.loc[0, 'Utilizator']
        r[6].loc[nr6, 'Observaţii'] = entry.loc[0, 'Obs']
        r[6].to_csv('registre/6.csv', index=False, header=True)

    # Centrifuga Sigma 2-16P
    if entry.loc[0, 'aparate'][7]:
        nr7 = len(r[7])
        r[7].loc[nr7, 'Data'] = entry.loc[0, 'Data_in']
        r[7].loc[nr7, 'Denumire probă'] = entry.loc[0, 'Produs']
        r[7].loc[nr7, 'Tipul probei/Lot'] = f'{entry.loc[0, "Tip_produs"]}/ {entry.loc[0, "Seria"]}'
        r[7].loc[nr7, 'Timp de centrifugare'] = entry.loc[0, 'aparate1'][7]
        r[7].loc[nr7, 'RPM'] = entry.loc[0, 'aparate2'][7]
        r[7].loc[nr7, 'Semnătura operator'] = entry.loc[0, 'Utilizator']
        r[7].loc[nr7, 'Curățarea/Statusul aparatului'] = entry.loc[0, 'stare'][7]
        r[7].loc[nr7, 'Observaţii'] = entry.loc[0, 'Obs']
        r[7].to_csv('registre/7.csv', index=False, header=True)

    # Etuva Memmert UN 110
    if entry.loc[0, 'aparate'][8]:
        nr8 = len(r[8])
        r[8].loc[nr8, 'Nr. crt.'] = nr8 + 1
        r[8].loc[nr8, 'Data'] = entry.loc[0, 'Data_in']
        r[8].loc[nr8, 'Tipul operaţiei'] = entry.loc[0, 'Operatiune']
        r[8].loc[nr8, 'Temperatura de operare (°C)'] = entry.loc[0, 'aparate1'][8]
        r[8].loc[nr8, 'Denumire produs'] = entry.loc[0, "Produs"]
        r[8].loc[nr8, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[8].loc[nr8, 'Semnătura persoanei care a efectuat  operaţia'] = entry.loc[0, 'Utilizator']
        r[8].loc[nr8, 'Observaţii'] = entry.loc[0, 'Obs']
        r[8].to_csv('registre/8.csv', index=False, header=True)

    # Sistem de dozare a proteinei Kjeldahl - Buchi
    if entry.loc[0, 'aparate'][9]:
        nr9 = len(r[9])
        r[9].loc[nr9, 'Nr. crt.'] = nr9 + 1
        r[9].loc[nr9, 'Data'] = entry.loc[0, 'Data_in']
        r[9].loc[nr9, 'Denumire operaţie'] = entry.loc[0, 'Operatiune']
        r[9].loc[nr9, 'Denumire produs, Seria, Data fabricației'] = \
            f'{entry.loc[0, "Produs"]}, {entry.loc[0, "Seria"]}, {entry.loc[0, "aparate1"][9]}'
        r[9].loc[nr9, 'Efectuat de/Semnătura'] = entry.loc[0, 'Utilizator']
        r[9].loc[nr9, 'Starea aparatului la finalul operării'] = entry.loc[0, 'stare'][9]
        r[9].to_csv('registre/9.csv', index=False, header=True)

    # Incubator Biosan ES - 20
    if entry.loc[0, 'aparate'][10]:
        nr10 = len(r[10])
        r[10].loc[nr10, 'Nr. crt.'] = nr10 + 1
        r[10].loc[nr10, 'Data'] = entry.loc[0, 'Data_in']
        r[10].loc[nr10, 'Operaţia'] = entry.loc[0, 'Operatiune']
        r[10].loc[nr10, 'Staus aparat'] = entry.loc[0, 'stare'][10]
        r[10].loc[nr10, 'Semnătura utilizator'] = entry.loc[0, 'Utilizator']
        r[10].loc[nr10, 'Observații'] = entry.loc[0, 'Obs']
        r[10].to_csv('registre/10.csv', index=False, header=True)

    # Liofilizator Cool Safe Model 55-4
    if entry.loc[0, 'aparate'][11]:
        nr11 = len(r[11])
        r[11].loc[nr11, 'Nr. crt.'] = nr11 + 1
        r[11].loc[nr11, 'Data'] = entry.loc[0, 'Data_in']
        r[11].loc[nr11, 'Denumire operaţie'] = entry.loc[0, 'Operatiune']
        r[11].loc[nr11, 'Denumire produs/Seria și data fabricației'] =\
            f'{entry.loc[0, "Produs"]}/ {entry.loc[0, "Seria"]}, {entry.loc[0, "aparate1"][11]}'
        r[11].loc[nr11, 'Efectuat operarea/Semnătura'] = entry.loc[0, 'Utilizator']
        r[11].loc[nr11, 'Starea aparatului la finalul operării'] = entry.loc[0, 'stare'][11]
        r[11].to_csv('registre/11.csv', index=False, header=True)

    # PH-metru Consort si Seven Copact S210
    if entry.loc[0, 'aparate'][12]:
        nr12 = len(r[12])
        r[12].loc[nr12, 'Data'] = entry.loc[0, 'Data_in']
        r[12].loc[nr12, 'Denumire produs'] = entry.loc[0, 'Produs']
        r[12].loc[nr12, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[12].loc[nr12, 'Tipul probei'] = entry.loc[0, 'Tip_produs']
        r[12].loc[nr12, 'pH determinat'] = entry.loc[0, 'aparate1'][12]
        r[12].loc[nr12, 'Efectuat determinarea si curățarea'] = entry.loc[0, 'Utilizator']
        r[12].loc[nr12, 'Starea aparatului la finalul operării'] = entry.loc[0, 'stare'][12]
        r[12].to_csv('registre/12.csv', index=False, header=True)

    # Evaporator Rotativ IKA RV 10 Control
    if entry.loc[0, 'aparate'][13]:
        nr13 = len(r[13])
        r[13].loc[nr13, 'Data'] = entry.loc[0, 'Data_in']
        r[13].loc[nr13, 'Denumire probă'] = entry.loc[0, 'Produs']
        r[13].loc[nr13, 'Tipul probei/Lot'] = f'{entry.loc[0, "Tip_produs"]}/ {entry.loc[0, "Seria"]}'
        r[13].loc[nr13, 'Timp de evaporare'] = entry.loc[0, 'aparate1'][13]
        r[13].loc[nr13, 'Semnătura operator'] = entry.loc[0, 'Utilizator']
        r[13].loc[nr13, 'Observaţii'] = entry.loc[0, 'Obs']
        r[13].to_csv('registre/13.csv', index=False, header=True)

    # Aparatul Soxhlet VELP
    if entry.loc[0, 'aparate'][14]:
        nr14 = len(r[14])
        r[14].loc[nr14, 'Nr. crt.'] = nr14 + 1
        r[14].loc[nr14, 'Data'] = entry.loc[0, 'Data_in']
        r[14].loc[nr14, 'Denumire probă/Produs'] = entry.loc[0, 'Produs']
        r[14].loc[nr14, 'Tipul probei/Lot'] = f'{entry.loc[0, "Tip_produs"]}/ {entry.loc[0, "Seria"]}'
        r[14].loc[nr14, 'Solvent utilizat'] = entry.loc[0, 'aparate1'][14]
        r[14].loc[nr14, 'Semnătura'] = entry.loc[0, 'Utilizator']
        r[14].loc[nr14, 'Observaţii'] = entry.loc[0, 'Obs']
        r[14].to_csv('registre/14.csv', index=False, header=True)

    # Vascozimetru Fungilab model Expert R
    if entry.loc[0, 'aparate'][15]:
        nr15 = len(r[15])
        r[15].loc[nr15, 'Data'] = entry.loc[0, 'Data_in']
        r[15].loc[nr15, 'Denumire probă/Lot'] = f'{entry.loc[0, "Produs"]}/ {entry.loc[0, "Seria"]}'
        r[15].loc[nr15, 'Vâscozitate (cP)'] = '-'
        r[15].loc[nr15, 'Timp de lucru'] = entry.loc[0, 'aparate1'][15]
        r[15].loc[nr15, 'Axul'] = entry.loc[0, 'aparate2'][15]
        r[15].loc[nr15, 'Semnătura operator'] = entry.loc[0, 'Utilizator']
        r[15].loc[nr15, 'Curățare/Statusul aparatului'] = entry.loc[0, 'stare'][15]
        r[15].loc[nr15, 'Observaţii'] = entry.loc[0, 'Obs']
        r[15].to_csv('registre/15.csv', index=False, header=True)

    # Extractor cu ultrasunete Steel 500-DG
    if entry.loc[0, 'aparate'][16]:
        nr16 = len(r[16])
        r[16].loc[nr16, 'Nr. crt.'] = nr16 + 1
        r[16].loc[nr16, 'Data'] = entry.loc[0, 'Data_in']
        r[16].loc[nr16, 'Operaţia/Denumirea probei'] = col_type
        r[16].loc[nr16, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[16].loc[nr16, 'Timp de extracție/Putere ultrasunete'] = entry.loc[0, 'aparate1'][16]
        r[16].loc[nr16, 'Nume și prenume operator'] = entry.loc[0, 'Utilizator']
        r[16].loc[nr16, 'Observaţii'] = entry.loc[0, 'Obs']
        r[16].to_csv('registre/16.csv', index=False, header=True)

    # Echipament determinare activitate apa Novasina
    if entry.loc[0, 'aparate'][17]:
        nr17 = len(r[17])
        r[17].loc[nr17, 'Nr. crt.'] = nr17 + 1
        r[17].loc[nr17, 'Data'] = entry.loc[0, 'Data_in']
        r[17].loc[nr17, 'Operaţia/Denumirea probei'] = col_type
        r[17].loc[nr17, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[17].loc[nr17, 'Valoarea activității apei cititie'] = entry.loc[0, 'aparate1'][17]
        r[17].loc[nr17, 'Nume și prenume operator'] = entry.loc[0, 'Utilizator']
        r[17].loc[nr17, 'Observaţii'] = entry.loc[0, 'Obs']
        r[17].to_csv('registre/17.csv', index=False, header=True)

    # Moara cu cutite Gm 200
    if entry.loc[0, 'aparate'][18]:
        nr18 = len(r[18])
        r[18].loc[nr18, 'Nr. crt.'] = nr18 + 1
        r[18].loc[nr18, 'Data'] = entry.loc[0, 'Data_in']
        r[18].loc[nr18, 'Operaţia/Denumirea probei'] = col_type
        r[18].loc[nr18, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[18].loc[nr18, 'Nume și prenume operator'] = entry.loc[0, 'Utilizator']
        r[18].loc[nr18, 'Observaţii'] = entry.loc[0, 'Obs']
        r[18].to_csv('registre/18.csv', index=False, header=True)

    # HPLC
    if entry.loc[0, 'aparate'][19]:
        nr19 = len(r[19])
        r[19].loc[nr19, 'Nr. crt.'] = nr19 + 1
        r[19].loc[nr19, 'Data'] = entry.loc[0, 'Data_in']
        r[19].loc[nr19, 'Denumire produs'] = entry.loc[0, "Produs"]
        r[19].loc[nr19, 'Seria și data fabricației'] = f'{entry.loc[0, "Seria"]}, {entry.loc[0, "aparate1"][19]}'
        r[19].loc[nr19, 'Efectuat determinarea'] = entry.loc[0, 'Utilizator']
        r[19].loc[nr19, 'Starea aparatului la finalul operării'] = entry.loc[0, 'stare'][19]
        r[19].to_csv('registre/19.csv', index=False, header=True)

    # UV VIS
    if entry.loc[0, 'aparate'][20]:
        nr20 = len(r[20])
        r[20].loc[nr20, 'Nr. crt.'] = nr20 + 1
        r[20].loc[nr20, 'Data'] = entry.loc[0, 'Data_in']
        r[20].loc[nr20, 'Denumire produs'] = entry.loc[0, "Produs"]
        r[20].loc[nr20, 'Serie/Lot'] = entry.loc[0, 'Seria']
        r[20].loc[nr20, 'Tipul probei'] = entry.loc[0, "Tip_produs"]
        r[20].loc[nr20, 'Absorbanța/Lungimea de undă'] = entry.loc[0, "aparate1"][20]
        r[20].loc[nr20, 'Efectuat citirea'] = entry.loc[0, 'Utilizator']
        r[20].loc[nr20, 'Starea aparatului la finalul operării'] = entry.loc[0, 'stare'][20]
        r[20].to_csv('registre/20.csv', index=False, header=True)

    # Probe
    if entry.loc[0, 'Operatiune'] == 'Analiză':
        nr21 = len(r[21])
        r[21].loc[nr21, 'Nr.'] = nr21 + 1
        r[21].loc[nr21, 'Denumire probă'] = entry.loc[0, "Produs"]
        r[21].loc[nr21, 'Seria'] = entry.loc[0, 'Seria']
        r[21].loc[nr21, 'Solicitant'] = entry.loc[0, 'Solicitant']
        r[21].loc[nr21, 'Analiza solicitată'] = entry.loc[0, "Analiza"]
        r[21].loc[nr21, 'Data primirii'] = entry.loc[0, 'Data_in']
        r[21].loc[nr21, 'Data eliberării'] = entry.loc[0, 'Data_elib']
        r[21].loc[nr21, 'Nume'] = entry.loc[0, 'Utilizator']
        r[21].loc[nr21, 'Semnătura'] = entry.loc[0, 'Utilizator']
        r[21].to_csv('registre/21.csv', index=False, header=True)

    # Registru rapoarte
    if entry.loc[0, 'Operatiune'] == 'Analiză':
        nr22 = len(r[22])
        r[22].loc[nr22, 'Nr. raport de analiză'] = nr22 + 1
        r[22].loc[nr22, 'Data primirii probei'] = entry.loc[0, 'Data_in']
        r[22].loc[nr22, 'Data eliberării raportului de analiză'] = entry.loc[0, 'Data_elib']
        r[22].loc[nr22, 'Denumire produs'] = entry.loc[0, 'Produs']
        r[22].loc[nr22, 'Serie'] = entry.loc[0, 'Seria']
        r[22].loc[nr22, 'Rezoluție'] = entry.loc[0, 'Rezultat']
        r[22].loc[nr22, 'Semnătura'] = entry.loc[0, 'Utilizator']
        r[22].to_csv('registre/22.csv', index=False, header=True)


def get_entry_values():

    # Get values from form fields
    operatiune = combo_operatiune.get()
    nume_proba = entry_nume_proba.get()
    tip_proba = combo_tip_proba.get()
    seria = entry_seria.get()
    analiza = entry_analiza.get()
    solicitant = entry_solicitant.get()
    data_intrare = entry_data_intrare.get_date().strftime('%d-%m-%Y')
    data_eliberare = entry_data_eliberare.get_date().strftime('%d-%m-%Y')
    utilizator = entry_utilizator.get()
    rezultat = entry_rezultat.get()
    observatii = entry_observatii.get()

    chx = [checkbox_val[i].get() for i in range(21)]
    ap1 = [param_ap1[i].get() for i in range(21)]
    ap2 = [param_ap2[i].get() for i in range(21)]
    st = [1] * 21

    # Clear form fields after submission
    combo_operatiune.current(0)
    entry_nume_proba.delete(0, tk.END)
    combo_tip_proba.current(0)
    entry_seria.delete(0, tk.END)
    entry_analiza.delete(0, tk.END)
    entry_data_intrare.set_date(datetime.now())
    entry_data_eliberare.set_date(datetime.now())
    entry_utilizator.delete(0, tk.END)
    entry_rezultat.delete(0, tk.END)
    entry_observatii.delete(0, tk.END)

    g = pd.read_csv('inregistrari.csv')

    entry_list = [len(g) + 1, operatiune, data_intrare, data_eliberare, nume_proba, tip_proba,
                  seria, analiza, solicitant, utilizator, rezultat, observatii, chx, ap1, ap2, st]

    entry_df = pd.DataFrame(entry_list).T
    entry_df.columns = g.columns
    entry_df = entry_df.replace('', '-')
    print(entry_df.Seria[0])

    g = pd.concat([g, entry_df], axis=0)
    g.to_csv('inregistrari.csv', index=False, header=True)

    completare_registre(entry_df)


def show(ii):

    df = pd.read_csv(f'registre/{ii}.csv')

    # Create the main window
    sh = tk.Tk()
    sh.title(f"Registru {df.columns[2]}")

    header_frame = ttk.Frame(sh, padding=20)
    header_frame.grid(row=0, column=0, sticky='ew')
    header_text = 'S.C. HOFIGAL EXPORT IMPORT S.A.\nDepartament Cercetare Dezvoltare Brevete'
    label_header = ttk.Label(header_frame, text=header_text, font=('Terminal', '12', 'bold'), cursor='heart')
    label_header.pack(padx=10, pady=15, anchor='w')
    label_title = ttk.Label(header_frame, text=f'{df.columns[1]}\n{df.columns[2]}',
                            font=('Arial', '12', 'bold'), justify='center')
    label_title.pack(padx=10, pady=15, anchor='center')

    table_frame = ttk.Frame(sh, padding=20)
    table_frame.grid(row=1, column=0)

    footer_frame = ttk.Frame(sh, padding=20)
    footer_frame.grid(row=2, column=0, sticky='ew')
    label_format = ttk.Label(footer_frame, text=f"Format - {int(float(df.columns[0]))}",
                             font=('Terminal', '10'))
    label_format.pack(padx=10, pady=5, anchor='w')
    label_pagina = ttk.Label(footer_frame, text='Pagina: 1', font=('Arial', '10'))
    label_pagina.pack(padx=10, pady=5, anchor='e')

    # Create a Treeview widget

    tree = ttk.Treeview(table_frame, cursor='gumby', height=20)

    # Define the number of columns
    tree["columns"] = list(df.columns[3:])
    tree["show"] = "headings"  # remove the first empty column

    # Create the column headers
    for column in df.columns[3:]:
        tree.heading(column, text=column)

    for col in df.columns[3:]:
        tree.column(col, width=150, minwidth=150, anchor='center', stretch=True)

    # Add the data rows
    for index, row in df.iterrows():
        tree.insert("", "end", values=list(row[3:]))

    # Pack the Treeview widget into the window
    tree.grid(row=0, column=0, sticky="nsew")

    # Create a vertical scrollbar
    vertical_scrollbar = tk.Scrollbar(table_frame, orient="vertical", command=tree.yview)
    vertical_scrollbar.grid(row=0, column=1, sticky="ns")

    # Create a horizontal scrollbar
    horizontal_scrollbar = tk.Scrollbar(table_frame, orient="horizontal", command=tree.xview)
    horizontal_scrollbar.grid(row=1, column=0, sticky="ew")

    # Link vertical scrollbar to the text widget
    tree.config(yscrollcommand=vertical_scrollbar.set)

    # Link horizontal scrollbar to the text widget
    tree.config(xscrollcommand=horizontal_scrollbar.set)

    # Start the Tkinter main loop
    sh.mainloop()


def toword(ii):
    # Read CSV file into a DataFrame
    csv_file = f'registre/{ii}.csv'
    df = pd.read_csv(csv_file)
    df2 = df.iloc[:, 3:]
    print(df2)

    # Create a new Word document
    doc = Document()

    # Add header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = 'S.C. HOFIGAL EXPORT IMPORT S.A.\nDepartament Cercetare Dezvoltare Brevete'

    # Add a title or any text
    doc.add_heading(df.columns[1], level=1)
    doc.add_paragraph(df.columns[2])

    # Add a table to the Word document
    table = doc.add_table(rows=1, cols=len(df2.columns))
    table.style = 'Table Grid'

    # Add the header row
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df2.columns):
        hdr_cells[i].text = column_name

    # Add the rest of the data
    for index, row in df2.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)

    # Add a footer with page number
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f"Format: {int(float(df.columns[0]))}                                            " \
                            f"                                                                        Pagina "

    # To add page number after the footer text
    run = footer_paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Add a footer
    footer_paragraph.add_run(" din ")

    # Add total pages
    run = footer_paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Save the document
    doc_file = f'registre-word/{df.columns[2]}.docx'
    doc.save(doc_file)

    print(f"Document saved as {doc_file}")


root = tk.Tk()
root.title('form')

# <editor-fold desc="frame1">

frame1 = ttk.Frame(root, padding='10')
frame1.grid(row=0, column=0)
row_nr = 0

label_operatiune = ttk.Label(frame1, text="Tip operațiune:")
label_operatiune.grid(row=row_nr, column=0, padx=10, pady=5)
options1 = ["Analiză", "Calibrare", "Verificare", "Curățare"]
combo_operatiune = ttk.Combobox(frame1, values=options1, state='readonly', justify='center', width=27)
combo_operatiune.current(0)
combo_operatiune.grid(row=row_nr, column=1, padx=10, pady=5)
combo_operatiune.bind("<<ComboboxSelected>>", disable_entries)
row_nr += 1

label_nume_proba = ttk.Label(frame1, text="Nume produs")
label_nume_proba.grid(row=row_nr, column=0, padx=10, pady=5)
entry_nume_proba = ttk.Entry(frame1, width=30)
entry_nume_proba.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1

label_tip_proba = ttk.Label(frame1, text="Tip produs:")
label_tip_proba.grid(row=row_nr, column=0, padx=10, pady=5)
options2 = ["-", "Materie prima", "Vrac", "Produs finit", "Cercetare", "Producție", "Diverse"]
combo_tip_proba = ttk.Combobox(frame1, values=options2, state='readonly', justify='center', width=27)
combo_tip_proba.current(0)
combo_tip_proba.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1

label_seria = ttk.Label(frame1, text="Seria:")
label_seria.grid(row=row_nr, column=0, padx=10, pady=5)
entry_seria = ttk.Entry(frame1, width=30)
entry_seria.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1

label_analiza = ttk.Label(frame1, text="Analiza:")
label_analiza.grid(row=row_nr, column=0, padx=10, pady=5)
entry_analiza = ttk.Entry(frame1, width=30)
entry_analiza.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1

label_solicitant = ttk.Label(frame1, text="Solicitant:")
label_solicitant.grid(row=row_nr, column=0, padx=10, pady=5)
entry_solicitant = tk.StringVar()
entry_solicitantCD = ttk.Radiobutton(frame1, text="CD", variable=entry_solicitant, value="CD")
entry_solicitantCC = ttk.Radiobutton(frame1, text="CC", variable=entry_solicitant, value="CC")
entry_solicitantCD.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1
entry_solicitantCC.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1

label_data_intrare = ttk.Label(frame1, text='Data intrare')
label_data_intrare.grid(row=row_nr, column=0, padx=10, pady=5)
entry_data_intrare = DateEntry(frame1, width=12, background='darkblue', foreground='white', borderwidth=2)
entry_data_intrare.grid(row=row_nr, column=1, padx=10, pady=5)
entry_data_intrare.set_date(datetime.now())
row_nr += 1

label_data_eliberare = ttk.Label(frame1, text='Data eliberare')
label_data_eliberare.grid(row=row_nr, column=0, padx=10, pady=5)
entry_data_eliberare = DateEntry(frame1, width=12, background='darkblue', foreground='white', borderwidth=2)
entry_data_eliberare.grid(row=row_nr, column=1, padx=10, pady=5)
entry_data_eliberare.set_date(datetime.now())
row_nr += 1

label_utilizator = ttk.Label(frame1, text="Utilizator:")
label_utilizator.grid(row=row_nr, column=0, padx=10, pady=5)
entry_utilizator = ttk.Entry(frame1, width=30)
entry_utilizator.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1

label_rezultat = ttk.Label(frame1, text="Rezultat:")
label_rezultat.grid(row=row_nr, column=0, padx=10, pady=5)
entry_rezultat = ttk.Entry(frame1, width=30)
entry_rezultat.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1

label_observatii = ttk.Label(frame1, text="Observații:")
label_observatii.grid(row=row_nr, column=0, padx=10, pady=5)
entry_observatii = ttk.Entry(frame1, width=30)
entry_observatii.grid(row=row_nr, column=1, padx=10, pady=5)
row_nr += 1

get_values_button = ttk.Button(frame1, text="Adaugă intrare", command=get_entry_values)
get_values_button.grid(row=row_nr, column=0, columnspan=2, padx=3, pady=3)
row_nr += 1

# </editor-fold>

# <editor-fold desc="frame2">

frame2 = ttk.Frame(root, padding="10")
frame2.grid(row=0, column=1, sticky='nsew')

checkbox_val = [tk.BooleanVar() for i in range(21)]
checkbox = [ttk.Checkbutton(frame2, text=f.aparat[i], variable=checkbox_val[i],
                            command=lambda i=i: toggle(i)) for i in range(21)]

label_ap1 = [ttk.Label(frame2, text=f.val1[i]) for i in range(21)]
label_ap2 = [ttk.Label(frame2, text=f.val2[i]) for i in range(21)]

param_ap1 = [ttk.Entry(frame2, width=20, state='disabled') for i in range(21)]
param_ap2 = [ttk.Entry(frame2, width=20, state='disabled') for i in range(21)]

butoane_word = [ttk.Button(frame2, text='Word', command=lambda j=j: toword(j)) for j in range(21)]

for i in range(21):
    checkbox[i].grid(row=i, column=0, sticky=tk.W)

    label_ap1[i].grid(row=i, column=1)
    param_ap1[i].grid(row=i, column=2, sticky='ew')

    if f.val1[i] == '-':
        label_ap1[i].grid_remove()
        param_ap1[i].grid_remove()

    label_ap2[i].grid(row=i, column=3)
    param_ap2[i].grid(row=i, column=4, sticky='ew')

    if f.val2[i] == '-':
        label_ap2[i].grid_remove()
        param_ap2[i].grid_remove()

    butoane_word[i].grid(row=i, column=5, sticky='w', padx=10)

ttk.Button(frame2, text='Registru Probe', command=lambda i=i: toword(21)).grid(row=21, column=4, sticky='w')
ttk.Button(frame2, text='Registru Rapoarte', command=lambda i=i: toword(22)).grid(row=21, column=5, sticky='w')

# </editor-fold>

# <editor-fold desc="frame3">
frame3 = ttk.Frame(root, padding="10")
frame3.grid(row=1, column=0, columnspan=2, sticky='nsew')

butoane_registre = [ttk.Button(frame3, text=f.aparat[i], command=lambda i=i: show(i)).grid(
    row=i//5, column=i % 5, pady=10) for i in range(23)]

# </editor-fold>

root.mainloop()
