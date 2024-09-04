import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


def toword(ii):

    # Read CSV file into a DataFrame
    csv_file = f'registre/{ii}.csv'
    df = pd.read_csv(csv_file)
    df2 = df.iloc[:, 3:]
    print(df2)

    # Create a new Word document
    doc = Document()

    #Add header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = 'S.C. HOFIGAL EXPORT IMPORT S.A.\nDepartament Cercetare Dezvoltare Brevete'

    # Add a title or any text
    doc.add_heading(df.columns[1], level=1)
    doc.add_paragraph(df.columns[2])

    # Add a table to the Word document
    table = doc.add_table(rows=1, cols=len(df2.columns))
    table.style = 'Table Grid'

    # Add the header row
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df2.columns):
        hdr_cells[i].text = column_name

    # Add the rest of the data
    for index, row in df2.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)

    # Add a footer with page number
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f"Format:{df.columns(0)}                                          " \
                            f"                                Page "

    # To add page number after the footer text
    run = footer_paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Add a footer
    footer_paragraph.add_run(" of ")

    # Add total pages
    run = footer_paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Save the document
    doc_file = f'registre-word/{df.columns[2]}.docx'
    doc.save(doc_file)

    print(f"Document saved as {doc_file}")
